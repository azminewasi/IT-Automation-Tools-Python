import pandas as pd 
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import sys

sheet_name="Model Test by Mahir Dyan Amin"
test_no=5

form = pd.read_excel("Course Deployment _ 44th BCS Final Model Test.xlsx",sheet_name=sheet_name)
form.fillna(0)


document = Document("template.docx")

print(form.info())
yn=input("Proceed?(y/n)\n-- ")


if yn=="n":
    sys.exit()
    

for i in range(200):
    p = document.add_paragraph("")
    p.add_run(str(i+1)+". "+str(form['Question'][i])+"\n").bold = True
    
    text="A. "+str(form['Option 1'][i])+"\n"+"B. "+str(form['Option 2'][i])+"\n"+"C. "+str(form['Option 3'][i])+"\n"+"D. "+str(form['Option 4'][i])+"\n"
    p.add_run(text)
    p.add_run("Answer: ").bold = True
    
    if form['Answer'][i]==1 or form['Answer'][i]=="১" :
        ansTxt=form['Option 1'][i]
    elif form['Answer'][i]==2 or form['Answer'][i]=="২":
        ansTxt=form['Option 2'][i]
    elif form['Answer'][i]==3 or form['Answer'][i]=="৩":
        ansTxt=form['Option 3'][i]
    elif form['Answer'][i]==4 or form['Answer'][i]=="৪":
        ansTxt=form['Option 4'][i]
    
    wp = p.add_run(str(ansTxt))
    wp.font.color.rgb = RGBColor(0,176,80)
    
    if str(form['Explanation'][i])=="nan":
        continue
    else:
        p.add_run("\n\nExplanation: ").bold = True
        p.add_run(str(form['Explanation'][i])+"\n")


document.save(str(test_no)+sheet_name+'.docx')
